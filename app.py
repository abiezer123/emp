import os
from flask import Flask, render_template, request, jsonify, redirect, url_for, session
from pymongo import MongoClient
from bson.objectid import ObjectId
from datetime import datetime, timedelta
from dotenv import load_dotenv
import re
from flask import send_file
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from calendar import monthrange
from bson.json_util import dumps
import io
from collections import Counter
import cloudinary
import cloudinary.uploader
import cloudinary.api
from datetime import date, datetime
from pymongo import UpdateOne
import calendar

load_dotenv()  # Load .env file

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'your_secret_key')  # Set a secret key for session management

MONGO_URI = os.getenv('MONGO_URI')
client = MongoClient(MONGO_URI)
db = client.attendance_db
attendance_collection = db.attendance
users_collection = db.users  # Assuming you have a users collectionz
members_collection = db.members


cloudinary.config( 
  cloud_name = "dwtixju79", 
  api_key = "836289374764364", 
  api_secret = "wWPaBCiM928AH2AczGjoJGG-Mxk")

def init_admin():
    if users_collection.count_documents({'username': 'admin'}) == 0:
        users_collection.insert_one({'username': 'admin', 'password': 'empire123'})

with app.app_context():
    init_admin()
    
@app.route('/')
def index():
    return redirect(url_for('login'))  # Redirect to login page

@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        username = request.form.get('username').strip()
        password = request.form.get('password')

        user = users_collection.find_one({'username': username})
        if user and user.get('password') == password:
            session['username'] = username
            return redirect(url_for('attendance'))  # Redirect to the attendance page
        else:
            error = 'Invalid username or password'
    return render_template('login.html', error=error)

@app.route('/attendance')
def attendance():
    if 'username' not in session:
        return redirect(url_for('login'))
    return render_template('attendance.html')

# add visitor 
@app.route('/api/attendance', methods=['POST'])
def add_attendance():
    try:
        data = request.get_json(force=True)
        name = data.get('name', '').strip()
        date_str = data.get('date', '').strip()
        is_visitor = bool(data.get('is_visitor', False))
        is_online = bool(data.get('is_online', False))
        is_child = bool(data.get('is_child', False))

        if not name or not date_str:
            return jsonify({'error': 'Name and date are required'}), 400

        # Parse date from string
        try:
            date = datetime.fromisoformat(date_str)
        except ValueError:
            return jsonify({'error': 'Invalid date format'}), 400

        # Normalize to day boundaries
        day_start = datetime(date.year, date.month, date.day)
        day_end = day_start + timedelta(days=1)

        # Ensure uniqueness: same name + same day
        existing = attendance_collection.find_one({
            "name": name,
            "date": {"$gte": day_start, "$lt": day_end}
        })

        if existing:
            return jsonify({'error': f'{name} already has attendance for this date'}), 409

        # Insert with Python datetime (Mongo stores as BSON Date)
        record = {
            'name': name,
            'date': date,
            'is_visitor': is_visitor,
            'is_online': is_online,
            'is_child': is_child,

        }
        attendance_collection.insert_one(record)

        return jsonify({'message': 'Attendance added successfully'}), 201

    except Exception as e:
        print("Error in add_attendance:", e)  # Debug log
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/attendance', methods=['GET'])
def get_attendance():
    date_str = request.args.get('date')
    if not date_str:
        return jsonify({'error': 'Date query parameter is required'}), 400
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
    except ValueError:
        return jsonify({'error': 'Invalid date format, expected YYYY-MM-DD'}), 400

    start = datetime(date_obj.year, date_obj.month, date_obj.day)
    end = datetime(date_obj.year, date_obj.month, date_obj.day, 23, 59, 59, 999999)

    cursor = attendance_collection.find({'date': {'$gte': start, '$lte': end}})
    records = []
    for doc in cursor:
        records.append({
            '_id': str(doc['_id']),
            'name': doc['name'],
            'date': doc['date'].isoformat(),
            'is_visitor': doc.get('is_visitor', False),
            'is_online': doc.get('is_online', False),
            'is_child': doc.get('is_child', False)
        })
    return jsonify({'records': records})

@app.route('/api/attendance/<record_id>', methods=['DELETE'])
def delete_attendance(record_id):
    try:
        result = attendance_collection.delete_one({'_id': ObjectId(record_id)})
        if result.deleted_count == 0:
            return jsonify({'error': 'Record not found'}), 404
        return jsonify({'message': 'Record deleted successfully'})
    except Exception:
        return jsonify({'error': 'Invalid record id'}), 400

@app.route('/api/names', methods=['GET'])
def suggest_names():
    q = request.args.get('q', '').strip()
    if not q:
        return jsonify([])  # Empty list if no query
    regex = re.compile(re.escape(q), re.IGNORECASE)
    names = attendance_collection.distinct('name', {'name': regex})[:10]
    return jsonify(names)

@app.route('/dashboard')
def dashboard():
    if 'username' not in session:
        return redirect(url_for('login'))
    return render_template('dashboard.html')

@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))

@app.route('/api/attendance_summary', methods=['GET'])
def get_attendance_summary():
    timeframe = request.args.get('timeframe', 'monthly')
    date_str = request.args.get('date')
    
    if not date_str:
        return jsonify({'error': 'Date query parameter is required'}), 400
    try:
        if timeframe == 'monthly':
            year, month = map(int, date_str.split('-'))
            last_day = monthrange(year, month)[1]
            start_date = datetime(year, month, 1)
            end_date = datetime(year, month, last_day, 23, 59, 59)
        else:  # yearly
            year = int(date_str)
            start_date = datetime(year, 1, 1)
            end_date = datetime(year, 12, 31, 23, 59, 59)
    except ValueError:
        return jsonify({'error': 'Invalid date format, expected YYYY-MM or YYYY'}), 400

    pipeline = [
        { '$match': { 'date': { '$gte': start_date, '$lte': end_date } } },
        { '$group': {
            '_id': { '$dateToString': { 'format': '%Y-%m-%d', 'date': '$date' } },
            'count': { '$sum': 1 }
        }},
        { '$sort': { '_id': 1 } }
    ]

    try:
        cursor = attendance_collection.aggregate(pipeline)
        summary_data = [{ 'date': doc['_id'], 'count': doc['count'] } for doc in cursor]
        return jsonify(summary_data)
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route('/api/download_attendance', methods=['GET'])
def download_attendance():
    year = int(request.args.get('year'))
    month = int(request.args.get('month'))

    start_date = datetime(year, month, 1)
    end_date = datetime(year + 1, 1, 1) if month == 12 else datetime(year, month + 1, 1)

    # Attendance Records
    records = list(attendance_collection.find({
        'date': {'$gte': start_date, '$lt': end_date}
    }))

    # Top Attendees
    top_attendees = list(attendance_collection.aggregate([
        { '$match': { 'date': {'$gte': start_date, '$lt': end_date} } },
        { '$group': { '_id': '$name', 'count': { '$sum': 1 } } },
        { '$sort': { 'count': -1 } },
        { '$limit': 10 }
    ]))

    # Attendance Summary (grouped by day)
    summary = list(attendance_collection.aggregate([
        { '$match': { 'date': {'$gte': start_date, '$lt': end_date} } },
        { '$group': {
            '_id': { '$dateToString': { 'format': "%Y-%m-%d", 'date': "$date" }},
            'count': { '$sum': 1 }
        }},
        { '$sort': { '_id': 1 }}
    ]))

    # Generate Word document
    doc = Document()
    doc.add_heading(f'Attendance Report - {start_date.strftime("%B %Y")}', 0)

    # Attendance Table
    doc.add_heading('Attendance Records', level=1)
    records_by_date = {}

    # Group records by date
    for record in records:
        date_obj = record['date']
        date_str = date_obj.strftime('%B %d, %Y')
        if date_str not in records_by_date:
            records_by_date[date_str] = []
        records_by_date[date_str].append(record)

    # Write grouped records
    for date_str in sorted(records_by_date):
        doc.add_paragraph(date_str, style='Heading 2')

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light List Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '#'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Time'

        for idx, record in enumerate(records_by_date[date_str], 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = record['name']
            row[2].text = record['date'].strftime('%I:%M %p')

        doc.add_paragraph()

    # Top Attendees
    doc.add_heading('Top Attendees', level=1)
    top_table = doc.add_table(rows=1, cols=2)
    top_table.style = 'Light List Accent 2'
    top_hdr = top_table.rows[0].cells
    top_hdr[0].text = 'Name'
    top_hdr[1].text = 'Attendance Count'
    for attendee in top_attendees:
        row = top_table.add_row().cells
        row[0].text = str(attendee['_id'])
        row[1].text = str(attendee['count'])

    # Attendance Summary
    doc.add_heading('Attendance Summary', level=1)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Light Grid Accent 3'
    hdr = summary_table.rows[0].cells
    hdr[0].text = 'Date'
    hdr[1].text = 'Attendance Count'
    summary_dates = []
    summary_counts = []
    for item in summary:
        row = summary_table.add_row().cells
        row[0].text = item['_id']
        row[1].text = str(item['count'])
        summary_dates.append(item['_id'])
        summary_counts.append(item['count'])

  
    # Return document
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    filename = f'attendance_{year}-{month:02d}.docx'
    return send_file(doc_stream, as_attachment=True, download_name=filename)

@app.route('/api/top_attendees', methods=['GET'])
def get_top_attendees():
    timeframe = request.args.get('timeframe', 'monthly')  # Default to monthly
    date_str = request.args.get('date')

    if not date_str:
        return jsonify({'error': 'Date query parameter is required'}), 400

    if timeframe == 'monthly':
        try:
            year, month = map(int, date_str.split('-'))
            start_date = datetime(year, month, 1)
            # Calculate the end date for the month
            if month == 12:
                end_date = datetime(year + 1, 1, 1)
            else:
                end_date = datetime(year, month + 1, 1)
        except ValueError:
            return jsonify({'error': 'Invalid date format for monthly summary, expected YYYY-MM'}), 400
    else:  # yearly
        try:
            year = int(date_str)
            start_date = datetime(year, 1, 1)
            end_date = datetime(year + 1, 1, 1)
        except ValueError:
            return jsonify({'error': 'Invalid date format for yearly summary, expected YYYY'}), 400

    pipeline = [
        {
            '$match': {
                'date': {'$gte': start_date, '$lt': end_date}
            }
        },
        {
            '$group': {
                '_id': '$name',
                'count': {'$sum': 1}
            }
        },
        {
            '$sort': {'count': -1}  # Sort by count descending
        }
    ]

    cursor = attendance_collection.aggregate(pipeline)
    top_attendees = []
    for doc in cursor:
        top_attendees.append({
            'name': doc['_id'],
            'count': doc['count']
        })

    return jsonify(top_attendees)

@app.route('/api/attendance/bulk-update', methods=['PUT'])
def bulk_update_attendance_dates():
    data = request.get_json()
    from_date_str = data.get('from_date')
    to_date_str = data.get('to_date')

    if not from_date_str or not to_date_str:
        return jsonify({'error': 'Both from_date and to_date are required'}), 400

    try:
        from_date_obj = datetime.fromisoformat(from_date_str)
        to_date_obj = datetime.fromisoformat(to_date_str)
    except ValueError:
        return jsonify({'error': 'Invalid date format'}), 400

    # Get all records on the same calendar day (regardless of time)
    day_start = from_date_obj.replace(hour=0, minute=0, second=0, microsecond=0)
    day_end = from_date_obj.replace(hour=23, minute=59, second=59, microsecond=999999)

    result = attendance_collection.update_many(
        {'date': {'$gte': day_start, '$lte': day_end}},
        {'$set': {'date': to_date_obj}}
    )

    return jsonify({'message': f'Updated {result.modified_count} record(s)'})



@app.route('/members')
def members():
    if 'username' not in session:
        return redirect(url_for('login'))

    # --- Step 1: Attendees Tab ---
    attendees = attendance_collection.distinct('name')
    attendees.sort()

    # --- Step 2: Find names with ≥2 occurrences in attendance ---
    pipeline = [
        {"$group": {"_id": "$name", "count": {"$sum": 1}}},
        {"$match": {"count": {"$gte": 2}}}
    ]
    member_docs = list(attendance_collection.aggregate(pipeline))
    member_names = [doc["_id"] for doc in member_docs]

    # --- Step 3: Insert only new names into members ---
    bulk_ops = []
    for name in member_names:
        bulk_ops.append(
            UpdateOne(
                {"name": name},  # filter by name
                {"$setOnInsert": {  # insert only if it does not exist
                    "name": name,
                    "birthdate": None,
                    "date_baptized": None,
                    "place_baptism": "",
                    "witnesses": "",
                    "father": "",
                    "mother": "",
                    "contact": "",
                    "email": "",
                    "facebook": "",
                    "address": "",
                    "image_url": ""
                }},
                upsert=True
            )
        )
    if bulk_ops:
        members_collection.bulk_write(bulk_ops)

    # --- Step 4: Fetch members for display ---
    members_cursor = members_collection.find({}, {"name": 1, "full_name": 1, "image_url": 1})
    members = [
        {
            "id": str(doc["_id"]),
            "name": doc.get("name", ""),
            "full_name": doc.get("full_name", ""),
            "image_url": doc.get("image_url", "")
        }
        for doc in members_cursor
    ]


    return render_template('members.html', attendees=attendees, members=members)

@app.route('/api/member_attendance/<name>', methods=['GET'])
def get_member_attendance(name):
    name = name.strip()
    if not name:
        return jsonify({'error': 'Name is required'}), 400

    # Get all unique attendance dates from the database
    all_dates_cursor = attendance_collection.find({}, {'date': 1})
    all_dates = set()
    for record in all_dates_cursor:
        all_dates.add(record['date'].date())  # Extract only the date part

    # Get all dates where the member is present
    present_dates_cursor = attendance_collection.find({'name': name}, {'date': 1})
    present_dates = set()
    for record in present_dates_cursor:
        present_dates.add(record['date'].date())

    # Prepare the response: every date in the database, marked present or not
    response = []
    for date in sorted(all_dates):
        response.append({
            'date': date.isoformat(),
            'present': date in present_dates
        })

    return jsonify(response)

# Helper to calculate age
def calculate_age(birthdate_str):
    if not birthdate_str:
        return None
    try:
        birthdate = datetime.strptime(birthdate_str, "%Y-%m-%d").date()
        today = date.today()
        return today.year - birthdate.year - ((today.month, today.day) < (birthdate.month, birthdate.day))
    except:
        return None

@app.route('/api/members/<member_id>', methods=['GET'])
def get_member(member_id):
    try:
        member = members_collection.find_one({"_id": ObjectId(member_id)})
        if not member:
            return jsonify({"error": "Member not found"}), 404

        member["_id"] = str(member["_id"])  # Convert ObjectId to string for JSON
        member["age"] = calculate_age(member.get("birthdate"))
        return jsonify(member)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/members/update", methods=["POST"])
def update_member():
    try:
        data = request.get_json()
        member_id = data.get("id")

        if not member_id:
            return jsonify({"success": False, "error": "Missing member ID"}), 400

        update_fields = {
            "full_name": data.get("full_name", "").strip(),
            "name": data.get("name", "").strip(),
            "birthdate": data.get("birthdate", "").strip(),
            "date_baptized": data.get("date_baptized", "").strip(),
            "place_baptism": data.get("place_baptism", "").strip(),
            "witnesses": data.get("witnesses", "").strip(),
            "father": data.get("father", "").strip(),
            "mother": data.get("mother", "").strip(),
            "contact": data.get("contact", "").strip(),
            "email": data.get("email", "").strip(),
            "facebook": data.get("facebook", "").strip(),
            "address": data.get("address", "").strip()
        }

        # Remove empty strings (to avoid overwriting with blank)
        update_fields = {k: v for k, v in update_fields.items() if v != ""}

        result = members_collection.update_one(
            {"_id": ObjectId(member_id)},
            {"$set": update_fields}
        )

        if result.matched_count == 0:
            return jsonify({"success": False, "error": "Member not found"}), 404

        return jsonify({"success": True, "message": "Member updated successfully!"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    
@app.route("/upload_image", methods=["POST"])
def upload_image():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]

    try:
        upload_result = cloudinary.uploader.upload(file)
        image_url = upload_result.get("secure_url")

        member_id = request.form.get("id")
        if member_id and image_url:
            members_collection.update_one(
                {"_id": ObjectId(member_id)},
                {"$set": {"image_url": image_url}},
                upsert=False
            )

        return jsonify({"url": image_url})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Jinja filter to convert month number to month name
@app.template_filter('month_name')
def month_name_filter(month_number):
    return calendar.month_name[month_number]  # 1 -> January, 2 -> February, etc.

@app.route('/birthdays')
def birthdays():
    members = list(members_collection.find())  # get all members

    months = [
        "January", "February", "March", "April", "May", "June",
        "July", "August", "September", "October", "November", "December"
    ]

    birthdays_dict = {month: [] for month in months}

    today = date.today()

    for m in members:
        if m.get('birthdate'):
            bdate = m['birthdate']  # 'YYYY-MM-DD'
            b = date.fromisoformat(bdate)
            month_name = months[b.month - 1]

            # Calculate current age
            current_age = today.year - b.year
            if (today.month, today.day) < (b.month, b.day):
                current_age -= 1
            next_age = current_age + 1

            birthdays_dict[month_name].append({
                'full_name': m.get('full_name') or m.get('name'),
                'birth_month': b.month,
                'birth_day': b.day,
                'current_age': current_age,
                'next_age': next_age
            })

    # Remove months with no birthdays
    birthdays_dict = {k: v for k, v in birthdays_dict.items() if v}

    # Sort each month’s list by day
    for month, members_list in birthdays_dict.items():
        birthdays_dict[month] = sorted(members_list, key=lambda x: x['birth_day'])

    return render_template("birthdays.html", birthdays=birthdays_dict, now=datetime.now())


if __name__ == '__main__':
    app.run(debug=True)